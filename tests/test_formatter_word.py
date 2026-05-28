"""
tests/test_formatter_word.py — Unit tests for WordFormatter.
"""

import logging
import pytest
from docx import Document
from docx.enum.section import WD_ORIENT

import sys, pathlib
sys.path.insert(0, str(pathlib.Path(__file__).parent.parent))

from formatter_word import WordFormatter


# ── fixtures ─────────────────────────────────────────────────────────────────

SAMPLE_ROWS = [
    ["ชื่อ", "แผนก", "ตำแหน่ง"],
    ["สมชาย ใจดี", "IT", "นักพัฒนา"],
    ["สมหญิง รักเรียน", "HR", "ผู้จัดการ"],
]


@pytest.fixture
def saved_docx(tmp_path):
    out = tmp_path / "test.docx"
    fmt = WordFormatter()
    fmt.write(SAMPLE_ROWS)
    fmt.save(out)
    return out


# ── file & structure ─────────────────────────────────────────────────────────

def test_file_created(saved_docx):
    assert saved_docx.exists()
    assert saved_docx.stat().st_size > 0


def test_file_readable(saved_docx):
    doc = Document(saved_docx)
    assert len(doc.tables) == 1


def test_row_count(saved_docx):
    doc = Document(saved_docx)
    assert len(doc.tables[0].rows) == len(SAMPLE_ROWS)


def test_column_count(saved_docx):
    doc = Document(saved_docx)
    assert len(doc.tables[0].columns) == 3


# ── cell values ──────────────────────────────────────────────────────────────

def test_header_values(saved_docx):
    doc = Document(saved_docx)
    texts = [c.text for c in doc.tables[0].rows[0].cells]
    assert texts == ["ชื่อ", "แผนก", "ตำแหน่ง"]


def test_data_row1_col0(saved_docx):
    doc = Document(saved_docx)
    assert doc.tables[0].rows[1].cells[0].text == "สมชาย ใจดี"


def test_data_row2_col1(saved_docx):
    doc = Document(saved_docx)
    assert doc.tables[0].rows[2].cells[1].text == "HR"


# ── formatting ───────────────────────────────────────────────────────────────

def test_header_bold(saved_docx):
    doc = Document(saved_docx)
    run = doc.tables[0].rows[0].cells[0].paragraphs[0].runs[0]
    assert run.bold is True


def test_data_not_bold(saved_docx):
    doc = Document(saved_docx)
    run = doc.tables[0].rows[1].cells[0].paragraphs[0].runs[0]
    assert not run.bold


def test_thai_font_on_header(saved_docx):
    doc = Document(saved_docx)
    run = doc.tables[0].rows[0].cells[0].paragraphs[0].runs[0]
    assert run.font.name == "TH Sarabun New"


def test_thai_font_on_data(saved_docx):
    doc = Document(saved_docx)
    run = doc.tables[0].rows[1].cells[0].paragraphs[0].runs[0]
    assert run.font.name == "TH Sarabun New"


def test_landscape_orientation(saved_docx):
    doc = Document(saved_docx)
    assert doc.sections[0].orientation == WD_ORIENT.LANDSCAPE


# ── multi-line cell ───────────────────────────────────────────────────────────

def test_multiline_cell_content(tmp_path):
    """\\n in a cell becomes a soft break — both lines appear in cell.text."""
    out = tmp_path / "ml.docx"
    fmt = WordFormatter()
    fmt.write([["Header"], ["บรรทัดแรก\nบรรทัดสอง"]])
    fmt.save(out)
    doc = Document(out)
    text = doc.tables[0].rows[1].cells[0].text
    assert "บรรทัดแรก" in text
    assert "บรรทัดสอง" in text


def test_multiline_stays_one_cell(tmp_path):
    """Multi-line value must not create extra rows in the table."""
    out = tmp_path / "ml2.docx"
    fmt = WordFormatter()
    fmt.write([["H"], ["a\nb\nc"]])
    fmt.save(out)
    doc = Document(out)
    assert len(doc.tables[0].rows) == 2   # header + 1 data row only


# ── edge cases ───────────────────────────────────────────────────────────────

def test_empty_rows_warning(caplog):
    fmt = WordFormatter()
    with caplog.at_level(logging.WARNING):
        fmt.write([])
    assert "No rows" in caplog.text


def test_single_row(tmp_path):
    """Header-only (single row) must not crash."""
    out = tmp_path / "single.docx"
    fmt = WordFormatter()
    fmt.write([["Col A", "Col B"]])
    fmt.save(out)
    assert Document(out).tables[0].rows[0].cells[0].text == "Col A"


def test_ragged_rows(tmp_path):
    """Rows shorter than header must be padded — no IndexError."""
    out = tmp_path / "ragged.docx"
    fmt = WordFormatter()
    fmt.write([["A", "B", "C"], ["only one"]])
    fmt.save(out)
    doc = Document(out)
    assert len(doc.tables[0].columns) == 3


def test_save_creates_parent_dirs(tmp_path):
    out = tmp_path / "sub" / "deep" / "out.docx"
    fmt = WordFormatter()
    fmt.write(SAMPLE_ROWS)
    fmt.save(out)
    assert out.exists()
