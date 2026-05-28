"""
formatter_word.py — Writes extracted rows to a formatted .docx file.
"""

import logging
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


_HEADER_BG   = "2E75B6"        # matches ExcelFormatter header colour
_THAI_FONT   = "TH Sarabun New"
_FONT_SIZE   = 11               # pt
_PAGE_W_CM   = 29.7             # A4 landscape width
_PAGE_H_CM   = 21.0             # A4 landscape height
_MARGIN_CM   = 1.5


class WordFormatter:
    """
    Writes a list of row-lists to a single Word table (.docx).

    Features:
    - A4 landscape page layout
    - Bold blue header row with white text  (mirrors ExcelFormatter style)
    - Thai font (TH Sarabun New) set on every text run via both the
      w:rFonts ascii and eastAsia/cs slots — prevents silent font fallback
    - Soft line-breaks for multi-line cells (Shift+Enter within one cell)
    - Equal column-width distribution across the printable area
    """

    def __init__(self, logger: logging.Logger | None = None) -> None:
        self.log = logger or logging.getLogger(__name__)
        self._doc = Document()
        self._setup_page()

    # ── public ───────────────────────────────────────────────────────────────

    def write(self, rows: list[list[str]]) -> None:
        """Append all rows as a Word table, then apply formatting."""
        if not rows:
            self.log.warning("No rows to write.")
            return

        num_cols = max(len(r) for r in rows)
        tbl = self._doc.add_table(rows=0, cols=num_cols)
        tbl.style = "Table Grid"   # built-in; present in every Word version
        self._set_col_widths(tbl, num_cols)

        for idx, row in enumerate(rows):
            tr = tbl.add_row()
            is_header = idx == 0
            padded = list(row) + [""] * (num_cols - len(row))
            for col_idx, value in enumerate(padded):
                self._fill_cell(tr.cells[col_idx], value or "", is_header)

        self.log.info("Formatted %d rows (%d columns).", len(rows), num_cols)

    def save(self, output_path: str | Path) -> None:
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self._doc.save(path)
        self.log.info("Saved: %s", path)

    # ── private ──────────────────────────────────────────────────────────────

    def _setup_page(self) -> None:
        section = self._doc.sections[0]
        section.orientation  = WD_ORIENT.LANDSCAPE
        section.page_width   = Cm(_PAGE_W_CM)
        section.page_height  = Cm(_PAGE_H_CM)
        section.left_margin  = Cm(_MARGIN_CM)
        section.right_margin = Cm(_MARGIN_CM)
        section.top_margin   = Cm(_MARGIN_CM)
        section.bottom_margin = Cm(_MARGIN_CM)

    def _set_col_widths(self, tbl, num_cols: int) -> None:
        """Distribute columns equally across the printable width."""
        available = Cm((_PAGE_W_CM - 2 * _MARGIN_CM) / num_cols)
        for col in tbl.columns:
            for cell in col.cells:
                cell.width = available

    def _fill_cell(self, cell, value: str, is_header: bool) -> None:
        if is_header:
            _set_cell_bg(cell, _HEADER_BG)

        p = cell.paragraphs[0]
        parts = value.split("\n") if "\n" in value else [value]

        for i, part in enumerate(parts):
            run = p.add_run(part)
            _apply_thai_font(
                run,
                bold=is_header,
                color_rgb=RGBColor(0xFF, 0xFF, 0xFF) if is_header else None,
            )
            # Add a soft return (Shift+Enter) between parts — stays in one cell
            if i < len(parts) - 1:
                run.add_break()


# ── module-level helpers ──────────────────────────────────────────────────────

def _apply_thai_font(
    run,
    bold: bool = False,
    color_rgb: RGBColor | None = None,
) -> None:
    """
    Apply Thai font to a run via all three w:rFonts slots.

    python-docx sets only w:rFonts/@w:ascii by default.  Thai characters
    are routed through the eastAsia and cs (complex-script) slots, so all
    three must be set to prevent Word from substituting a fallback font that
    may not render Thai correctly — especially on macOS.
    """
    run.bold = bold
    run.font.name = _THAI_FONT
    run.font.size = Pt(_FONT_SIZE)
    rPr   = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), _THAI_FONT)
    rFonts.set(qn("w:cs"),       _THAI_FONT)
    if color_rgb is not None:
        run.font.color.rgb = color_rgb


def _set_cell_bg(cell, hex_color: str) -> None:
    """
    Set cell background fill colour via raw Word XML.

    python-docx ≥ 1.x has no direct table-cell shading API; we write
    the <w:shd> element directly into <w:tcPr>.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)
